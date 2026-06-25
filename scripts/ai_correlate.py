#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI Correlator for SEC-CATm audits (Android-focused)

- Iterate through the requirements and decide: YES / NO / N/A / MANUAL.
- Avoid "YES" if there is no clear positive evidence; when in doubt → MANUAL.
- Optional preselection (src-catm*) based on capabilities detected in the code.
- Integrate evidence from SARIF (SAST), MobSF (static/dynamic), Trivy, and analysis of the AndroidManifest/network_security_config.
- Display progress [i/N] and percentage.
- Limit to N requirements with --max-requirements.
- Use OpenAI only if explicitly requested: --use-llm {never,uncertain,always} (default: never).
- Generate: audit-findings.json, secm-catm.docx, checklist.docx, audit-summary.docx/md.
- CI-friendly: on errors, print traceback and exit with code 0.
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import sys
import traceback
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Tuple, Optional
import time
import xml.etree.ElementTree as ET

OPENAI_AVAILABLE = False
try:
    from openai import OpenAI  # type: ignore
    OPENAI_AVAILABLE = True
except Exception:
    OPENAI_AVAILABLE = False

# ---- DOCX --------------------------------------------------------------------
DOCX_AVAILABLE = False
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ---- Constants --------------------------------------------------------------
DECISIONS = {"YES", "NO", "N/A", "MANUAL"}
CODE_EXTS = (".java", ".kt", ".xml", ".gradle", ".kts", ".properties", ".txt")
MAX_HINT_FILES = 800
MAX_HINT_MATCHES = 200

CAPABILITY_KEYWORDS = {
    "webview": ["webview", "setJavaScriptEnabled", "file://"],
    "biometric": ["BiometricPrompt", "USE_BIOMETRIC", "USE_FINGERPRINT"],
    "camera": ["android.permission.CAMERA"],
    "location": ["ACCESS_FINE_LOCATION", "ACCESS_COARSE_LOCATION"],
    "nfc": ["android.permission.NFC"],
    "ble": ["Bluetooth", "BLUETOOTH_CONNECT", "BLUETOOTH_SCAN"],
    "saml": ["SAML", "AuthnStatement", "Assertion"],
    "oauth": ["OAuth", "openid", "OpenID Connect", "AuthorizationCode"],
    "crypto": ["PBKDF2", "bcrypt", "scrypt", "MessageDigest", "AES", "Cipher"],
}

RISK_HINTS = [
    r"http://",
    r"android:debuggable=\"true\"",
    r"android:allowBackup=\"true\"",
    r"android:usesCleartextTraffic=\"true\"",
    r"MODE_WORLD_READABLE",
    r"MODE_WORLD_WRITEABLE",
    r"MD5",
    r"SHA-1",
    r"AES\/ECB",
    r"setJavaScriptEnabled\(true\)",
    r"cleartextTrafficPermitted=\"true\"",
]

POSITIVE_TLS_HINTS = [
    r"android:usesCleartextTraffic=\"false\"",
    r"NetworkSecurityConfig.*cleartextTrafficPermitted=\"false\"",
]

# ---- Utilities --------------------------------------------------------------
def read_text_safe(p: Path, limit: int | None = None) -> str:
    try:
        data = p.read_text(encoding="utf-8", errors="replace")
    except Exception:
        try:
            data = p.read_text(encoding="latin-1", errors="replace")
        except Exception:
            return ""
    return data if limit is None else data[:limit]


def load_json_safe(p: Path) -> Any:
    try:
        return json.loads(read_text_safe(p))
    except Exception:
        return None


def safe_json_extract(s: str) -> Dict[str, Any] | None:
    stack = 0
    start = None
    for i, ch in enumerate(s):
        if ch == "{":
            if stack == 0:
                start = i
            stack += 1
        elif ch == "}":
            stack -= 1
            if stack == 0 and start is not None:
                frag = s[start : i + 1]
                try:
                    return json.loads(frag)
                except Exception:
                    pass
    return None


def load_checklist(path: Path) -> List[Dict[str, Any]]:
    raw = load_json_safe(path)
    if raw is None:
        raise ValueError(f"Invalid JSON: {path}")
    if isinstance(raw, dict):
        items = raw.get("items") or raw.get("requirements") or raw.get("data") or []
    elif isinstance(raw, list):
        items = raw
    else:
        items = []
    out: List[Dict[str, Any]] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        puid = it.get("PUID") or it.get("puid") or it.get("id") or ""
        it["PUID"] = str(puid).strip()
        out.append(it)
    return out


# ---- Quick code hints & capabilities -----------------------------------------
@dataclass
class CodeHints:
    matches: List[str]    # "file:line: snippet"
    present: set[str]     # present risk hints (regex.pattern)
    caps: set[str]        # inferred capabilities


def quick_code_hints(root: Path) -> CodeHints:
    matches: List[str] = []
    present: set[str] = set()
    caps: set[str] = set()
    files: List[Path] = []

    for ext in CODE_EXTS:
        files.extend([p for p in root.rglob(f"*{ext}")][:MAX_HINT_FILES])

    pat_list = [re.compile(pat, re.IGNORECASE) for pat in RISK_HINTS]

    for p in files:
        try:
            txt = read_text_safe(p, limit=100_000)
        except Exception:
            continue
        for pat in pat_list:
            for m in pat.finditer(txt):
                if len(matches) < MAX_HINT_MATCHES:
                    before = txt[: m.start()]
                    line_no = before.count("\n") + 1
                    snippet = txt[m.start() : m.start() + 120].replace("\n", " ")
                    matches.append(f"{p}:{line_no}: {snippet}")
                present.add(pat.pattern)

        low = txt.lower()
        for cap, kws in CAPABILITY_KEYWORDS.items():
            if any(k.lower() in low for k in kws):
                caps.add(cap)

    return CodeHints(matches=matches, present=present, caps=caps)


# ---- AndroidManifest & network_security_config -------------------------------
@dataclass
class ManifestInfo:
    manifest_path: Optional[Path]
    allow_backup: Optional[bool]
    debuggable: Optional[bool]
    uses_cleartext_traffic: Optional[bool]
    network_security_config_path: Optional[Path]
    n_exported_with_intentfilter: int
    exported_details: List[str]
    target_sdk: Optional[int]


def _bool_attr(val: Optional[str]) -> Optional[bool]:
    if val is None:
        return None
    v = val.strip().lower()
    if v in ("true", "1", "yes"):
        return True
    if v in ("false", "0", "no"):
        return False
    return None


def parse_network_security_config(app_root: Path, ns_path: Path) -> Dict[str, Any]:
    """Return {'base_cleartext_true': bool, 'any_domain_cleartext_true': bool}."""
    out = {"base_cleartext_true": False, "any_domain_cleartext_true": False}
    try:
        tree = ET.parse(ns_path)
        root = tree.getroot()
        # base-config
        for base in root.findall(".//base-config"):
            if base.get("cleartextTrafficPermitted", "").lower() == "true":
                out["base_cleartext_true"] = True
        # domain-config
        for d in root.findall(".//domain-config"):
            if d.get("cleartextTrafficPermitted", "").lower() == "true":
                out["any_domain_cleartext_true"] = True
    except Exception:
        pass
    return out


def parse_manifest(source_root: Path) -> ManifestInfo:
    manifest = None
    for p in source_root.rglob("AndroidManifest.xml"):
        manifest = p
        break

    allow_backup = debuggable = uses_cleartext = None
    n_exported = 0
    exported_details: List[str] = []
    target_sdk: Optional[int] = None
    ns_cfg_path: Optional[Path] = None

    if manifest and manifest.exists():
        try:
            tree = ET.parse(manifest)
            root = tree.getroot()
            # namespace
            AND = "{http://schemas.android.com/apk/res/android}"

            # targetSdkVersion if it exists in uses-sdk
            for uses in root.findall("uses-sdk"):
                v = uses.get(AND + "targetSdkVersion")
                if v and v.isdigit():
                    target_sdk = int(v)

            app = root.find("application")
            if app is not None:
                allow_backup = _bool_attr(app.get(AND + "allowBackup"))
                debuggable = _bool_attr(app.get(AND + "debuggable"))
                uses_cleartext = _bool_attr(app.get(AND + "usesCleartextTraffic"))
                # networkSecurityConfig
                ns_attr = app.get(AND + "networkSecurityConfig")
                if ns_attr and ns_attr.startswith("@xml/"):
                    xml_name = ns_attr.split("/", 1)[1] + ".xml"
                    # search in res/xml/**/xml_name
                    for candidate in source_root.rglob(f"**/{xml_name}"):
                        ns_cfg_path = candidate
                        break

            # components with intent-filter and exported=true
            for tag in ["activity", "service", "receiver", "provider"]:
                for comp in root.findall(f".//{tag}"):
                    exported = comp.get(AND + "exported")
                    has_intent = comp.find("intent-filter") is not None
                    if has_intent:
                        # If it has an intent-filter, an explicit exported attribute is mandatory on API 31+.
                        if exported is None:
                            exported_details.append(f"{tag} missing exported (intent-filter present)")
                        else:
                            val = _bool_attr(exported)
                            if val:
                                n_exported += 1
                                nm = comp.get(AND + "name") or "(anonymous)"
                                exported_details.append(f"{tag}:{nm} exported=true")
        except Exception:
            pass

    # add flags from the network_security_config
    if ns_cfg_path:
        ns = parse_network_security_config(source_root, ns_cfg_path)
        #If base/domain cleartext is true and uses_cleartext_traffic is still None → consider it a risk
        if ns.get("base_cleartext_true") or ns.get("any_domain_cleartext_true"):
            # it does not enforce uses_cleartext_traffic, but it is evidence of an insecure configuration
            pass

    return ManifestInfo(
        manifest_path=manifest,
        allow_backup=allow_backup,
        debuggable=debuggable,
        uses_cleartext_traffic=uses_cleartext,
        network_security_config_path=ns_cfg_path,
        n_exported_with_intentfilter=n_exported,
        exported_details=exported_details,
        target_sdk=target_sdk,
    )


# ---- Evidence parsing (SARIF, MobSF, Trivy) ------------------------------
@dataclass
class Evidence:
    sarif_rules: Dict[str, int]
    mobsf_hints: set[str]
    trivy_summary: Dict[str, Any]
    manifest: ManifestInfo
    ns_flags: Dict[str, Any]


def parse_sarif(p: Path) -> Dict[str, int]:
    data = load_json_safe(p)
    out: Dict[str, int] = {}
    if not isinstance(data, dict):
        return out
    for run in data.get("runs", []) or []:
        results = run.get("results", []) or []
        for r in results:
            rid = r.get("ruleId") or (r.get("rule") or {}).get("id") or "unknown"
            out[rid] = out.get(rid, 0) + 1
    return out


def mobsf_text_hints(json_file: Path) -> set[str]:
    s = read_text_safe(json_file, limit=2_000_000).lower()
    hints = set()
    for k in [
        "android:debuggable",
        "allowbackup",
        "webview",
        "http://",
        "https://",
        "aes/ecb",
        "md5",
        "cleartexttrafficpermitted",
        "setjavascriptenabled(true)",
        "usescleartexttraffic",
    ]:
        if k in s:
            hints.add(k)
    return hints


def collect_evidence(reports_dir: Path, source_root: Path) -> Evidence:
    sarif_rules: Dict[str, int] = {}
    mobsf_hints: set[str] = set()
    trivy_summary: Dict[str, Any] = {}

    for p in reports_dir.rglob("*.sarif"):
        try:
            rules = parse_sarif(p)
            for k, v in rules.items():
                sarif_rules[k] = sarif_rules.get(k, 0) + v
        except Exception:
            continue

    for p in reports_dir.rglob("mobsf*.json"):
        try:
            mobsf_hints |= mobsf_text_hints(p)
        except Exception:
            continue

    for p in reports_dir.rglob("*.json"):
        lname = p.name.lower()
        if "trivy" in lname or "agent_payload" in lname:
            data = load_json_safe(p)
            if isinstance(data, dict):
                trivy_summary["files"] = trivy_summary.get("files", []) + [p.name]
                sev_counts = trivy_summary.get("severity", {})
                payload = json.dumps(data)
                for key in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "UNKNOWN"]:
                    if key in payload:
                        sev_counts[key] = sev_counts.get(key, 0) + 1
                trivy_summary["severity"] = sev_counts

    # Manifest + network_security_config
    mani = parse_manifest(source_root)
    ns_flags = {}
    if mani.network_security_config_path:
        ns_flags = parse_network_security_config(source_root, mani.network_security_config_path)

    return Evidence(
        sarif_rules=sarif_rules,
        mobsf_hints=mobsf_hints,
        trivy_summary=trivy_summary,
        manifest=mani,
        ns_flags=ns_flags,
    )


# ---- Preselection src-catm* --------------------------------------------------
def requirement_seems_related(req: Dict[str, Any], caps: set[str]) -> bool:
    text = " ".join(str(req.get(k, "")) for k in ["PUID", "Requirement description", "Validation criteria", "Rationale", "Source"]).lower()
    conds = [
        ("saml", "saml"),
        ("webview", "webview"),
        ("biometric", "biometric"),
        ("bluetooth", "ble"),
        ("ble", "ble"),
        ("nfc", "nfc"),
        ("camera", "camera"),
        ("location", "location"),
        ("oauth", "oauth"),
    ]
    for needle, cap in conds:
        if needle in text and cap not in caps:
            return False
    return True


# ---- Context summary for the LLM ----------------------------------------
def summarize_reports_for_ai(ev: Evidence) -> str:
    lines = []
    sr = ", ".join(f"{k}:{v}" for k, v in sorted(ev.sarif_rules.items())[:30])
    lines.append(f"SARIF rules (top): {sr or 'none'}")
    lines.append(f"MobSF hints: {', '.join(sorted(ev.mobsf_hints)) or 'none'}")
    if ev.trivy_summary:
        sev = ev.trivy_summary.get("severity", {})
        lines.append(f"Trivy severity tags: {sev}")
    m = ev.manifest
    mline = f"Manifest: allowBackup={m.allow_backup} debuggable={m.debuggable} usesCleartextTraffic={m.uses_cleartext_traffic} exported(IF)={m.n_exported_with_intentfilter}"
    lines.append(mline)
    if ev.ns_flags:
        lines.append(f"NetSecConfig: {ev.ns_flags}")
    return "\n".join(lines)


# ---- MASTER PROMPT (long, with few-shot examples) -----------------------------------
def build_llm_prompt(puid: str, req: Dict[str, Any], app_ctx: str, code: 'CodeHints', ev: 'Evidence') -> str:
    desc = req.get("Requirement description") or req.get("description") or ""
    crit = req.get("Validation criteria") or req.get("criteria") or ""
    importance = req.get("Importance") or req.get("importance") or "Not described"

    code_hints_top = "; ".join(code.matches[:6]) or "—"
    code_hints_set = ", ".join(sorted(ev.manifest.exported_details + list(code.present))) or "none"
    sarif_top = ", ".join(f"{k}:{v}" for k, v in list(ev.sarif_rules.items())[:10]) or "none"
    mobsf_set = ", ".join(sorted(ev.mobsf_hints)) or "none"
    trivy_brief = json.dumps(ev.trivy_summary) if ev.trivy_summary else "None"

    return f"""
You are a **Senior Mobile Security Auditor** specialized in **Android healthcare apps**. Evaluate ONE requirement.

OUTPUT (STRICT JSON):
- "decision": "YES"|"NO"|"N/A"|"MANUAL"
- "evidence": 1–5 bullets (short)
- "rationale": 2–5 bullets (short, no chain-of-thought)
- "manual_steps": only if "MANUAL"

RULES:
- "N/A" if the capability clearly doesn't exist in app context.
- "NO" if there is contradictory evidence vs the requirement (e.g., allowBackup=true; debuggable=true; exported components w/ IF; cleartext allowed).
- "YES" only if requirement applies and there's positive evidence or absence of contradictions (be conservative).
- "MANUAL" when runtime behavior or evidence is insufficient/ambiguous.

FEW-SHOTS omitted for brevity (same as previous message).

CURRENT CASE
PUID: {puid}
Requirement description: {desc}
Validation criteria: {crit}
Importance: {importance}

App context:
{app_ctx.strip()[:2000]}

Evidence:
- Manifest/Exported: {code_hints_set}
- Code hints (examples): {code_hints_top}
- SARIF: {sarif_top}
- MobSF: {mobsf_set}
- Trivy: {trivy_brief}

Return STRICT JSON only.
""".strip()


# ---- OpenAI / Heuristics -----------------------------------------------------
def llm_decide(puid: str, req: Dict[str, Any], app_ctx: str, code: CodeHints, ev: Evidence, timeout_s: int) -> Tuple[str, str, Dict[str, Any]]:
    if not OPENAI_AVAILABLE or not os.getenv("LLM_API_KEY"):
        return ("N/A", "OpenAI not available", {})

    client = OpenAI(
        api_key=os.getenv("LLM_API_KEY"),
        base_url=os.getenv("LLM_BASE_URL"),
    )
    model = os.getenv("LLM_MODEL")
    max_out = int(os.getenv("LLM_MAX_OUTPUT_TOKENS", "900"))
    reasoning_effort = os.getenv("LLM_REASONING_EFFORT", "medium")

    prompt = build_llm_prompt(puid, req, app_ctx, code, ev)

    # Soft timeout per call (better than blocking the job)
    start = time.time()
    try:
        rsp = client.responses.create(
            model=model,
            input=prompt,
            reasoning={"effort": reasoning_effort},
            max_output_tokens=max_out,
        )
        if time.time() - start > timeout_s:
            return ("MANUAL", "LLM timeout -> manual review suggested", {"manual_steps": ["Revisar requisito en dispositivo/emulador y logs"]})
        text = getattr(rsp, "output_text", None)
        if not text:
            try:
                d = rsp.to_dict()
                choices = d.get("output", []) or d.get("choices", [])
                text = json.dumps(choices) if choices else ""
            except Exception:
                text = ""
    except Exception as e:
        return ("MANUAL", f"LLM error: {e}", {"manual_steps": ["Review manually due to insufficient context"]})

    obj = safe_json_extract(text or "")
    if not obj or "decision" not in obj:
        return ("MANUAL", "Unparseable LLM output", {"manual_steps": ["Review manually due to invalid output"]})

    decision = str(obj.get("decision", "MANUAL")).strip().upper()
    ev_field = obj.get("evidence", "")
    if isinstance(ev_field, list):
        evidence_text = " | ".join(str(x) for x in ev_field[:6])
    else:
        evidence_text = str(ev_field or "").strip()

    extras: Dict[str, Any] = {}
    if isinstance(obj.get("rationale"), list):
        extras["rationale"] = obj["rationale"][:6]
    if isinstance(obj.get("manual_steps"), list):
        extras["manual_steps"] = obj["manual_steps"][:8]

    if decision not in DECISIONS:
        decision = "MANUAL"
    return (decision, evidence_text or "LLM provided no evidence", extras)


def _req_has(text: str, words: List[str]) -> bool:
    low = text.lower()
    return any(w in low for w in words)


def heuristic_decide(req: Dict[str, Any], code: CodeHints, ev: Evidence) -> Tuple[str, str, bool]:
    """
    Conservative rules:

        - NO if there is direct contradictory evidence (e.g., allowBackup=true against a backup-off policy; debuggable=true in release; cleartext allowed when TLS is required; exported components with intent-filters without proper control)
        - N/A if the requirement depends on a missing capability (webview/biometric/NFC/etc.)
        - YES only with positive evidence or absence of contradiction + compliance hints
        - MANUAL if there is not enough evidence
    """
    desc = str(req.get("Requirement description", "") or req.get("description", ""))
    crit = str(req.get("Validation criteria", "") or req.get("criteria", ""))
    text = f"{desc}\n{crit}".lower()
    hints = code.present
    rules = ev.sarif_rules
    mani = ev.manifest
    ns = ev.ns_flags

    # 1) Specific rules per common requirement
    # TLS / Cleartext
    if _req_has(text, ["https", "tls", "cleartext", "http only forbidden", "no cleartext"]):
        cleartext = (mani.uses_cleartext_traffic is True) or ns.get("base_cleartext_true") or ns.get("any_domain_cleartext_true") or ("http://" in hints)
        if cleartext:
            return ("NO", "Cleartext permitted/used (Manifest or netsec config or code hints)", False)
        # positive evidence
        if (mani.uses_cleartext_traffic is False) or any(ph in hints for ph in POSITIVE_TLS_HINTS):
            return ("YES", "TLS enforced (no cleartext permitted)", False)
        # no contradiction but no strong evidence
        return ("MANUAL", "No visible contradiction; needs to be verified at runtime", True)

    # Backup policy
    if _req_has(text, ["backup", "allowbackup"]):
        if mani.allow_backup is True:
            return ("NO", "android:allowBackup=true", False)
        if mani.allow_backup is False:
            return ("YES", "android:allowBackup=false", False)
        return ("MANUAL", "allowBackup not declared -> requires review", True)

    # Debuggable
    if _req_has(text, ["debuggable", "release build debuggable", "no debuggable in prod"]):
        if mani.debuggable is True:
            return ("NO", "android:debuggable=true", False)
        if mani.debuggable is False:
            return ("YES", "android:debuggable=false", False)
        return ("MANUAL", "debuggable not declared -> revisar buildTypes", True)

    # Exported components
    if _req_has(text, ["export", "exported", "implicit exported", "component exposure"]):
        if mani.n_exported_with_intentfilter > 0:
            return ("NO", f"{mani.n_exported_with_intentfilter} component(s) exported with intent-filter", False)
        # If there is an intent-filter without an explicit exported attribute on target API 31+, it is a risk / not installable.
        if any("missing exported" in d for d in mani.exported_details):
            return ("NO", "intent-filter present without explicit exported (API31+)", False)
        return ("YES", "No exported components with IF", False)

    # Weak cryptography
    if _req_has(text, ["md5", "sha-1", "weak hash", "weak crypto", "ecb"]):
        if any(k for k in rules if re.search(r"(md5|sha-1|weak|ecb)", k, re.I)) or any("MD5" in m or "AES/ECB" in m for m in code.matches):
            return ("NO", "Weak crypto patterns/SARIF rules", False)
        return ("YES", "No weak crypto evidence", False)

    # PII logging
    if _req_has(text, ["pii", "logs", "no sensitive info in logs", "logging"]):
        if any("Log.d" in m or "Log.i" in m for m in code.matches):
            return ("MANUAL", "Needs to check logcat at runtime", True)
        return ("MANUAL", "No sufficient static evidence; runtime required", True)

    # Missing capabilities → N/A
    caps_text_pairs = [
        ("saml", "saml"),
        ("webview", "webview"),
        ("biometric", "biometric"),
        ("nfc", "nfc"),
        ("bluetooth", "ble"),
        ("ble", "ble"),
        ("camera", "camera"),
        ("location", "location"),
        ("oauth", "oauth"),
    ]
    for needle, cap in caps_text_pairs:
        if needle in text and cap not in code.caps:
            return ("N/A", f"Capability not present: {cap}", False)

    # General: if there are strong security-related indicators and no contradiction is found → MANUAL (conservative)
    if any(k in text for k in ["tls", "encrypt", "permission", "session", "token", "cookie", "audit", "pinning"]):
        if hints:
            return ("MANUAL", "Insufficient static evidence to conclude", True)
        return ("MANUAL", "Insufficient evidence", True)

    # By default
    return ("MANUAL", "Insufficient information to conclude", True)


# ---- Report generation --------------------------------------------------
def write_audit_findings(path: Path, results: Dict[str, Dict[str, Any]]) -> None:
    path.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")


def write_markdown_summary(path: Path, totals: Dict[str, int], notes: List[str]) -> None:
    md = []
    md.append("# Audit Summary\n")
    md.append(f"- Total requirements: {totals['total']}")
    md.append(f"- YES: {totals['YES']}")
    md.append(f"- NO: {totals['NO']}")
    md.append(f"- N/A: {totals['N/A']}")
    md.append(f"- MANUAL: {totals['MANUAL']}\n")
    if notes:
        md.append("## Notes\n")
        md.extend([f"- {n}" for n in notes])
    path.write_text("\n".join(md), encoding="utf-8")


def write_docx_checklist(path: Path, results: Dict[str, Dict[str, Any]], reqs_by_puid: Dict[str, Dict[str, Any]]) -> None:
    if not DOCX_AVAILABLE:
        return
    doc = Document()
    doc.add_heading("Checklist (SEC-CATm)", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "PUID"
    hdr[1].text = "Requirement"
    hdr[2].text = "Result"
    hdr[3].text = "Evidence"
    for puid, info in results.items():
        row = table.add_row().cells
        row[0].text = puid
        row[1].text = str(reqs_by_puid.get(puid, {}).get("Requirement description", ""))[:300]
        row[2].text = info.get("decision", "")
        row[3].text = str(info.get("evidence", ""))[:400]
    doc.save(path.as_posix())


def _append_requirement_block(doc, puid: str, info: Dict[str, Any], reqs_by_puid: Dict[str, Any]) -> None:
    doc.add_heading(puid, level=2)
    doc.add_paragraph(reqs_by_puid.get(puid, {}).get("Requirement description", ""))
    doc.add_paragraph(f"Result: {info.get('decision','')}")
    ev = info.get("evidence", "")
    if ev:
        doc.add_paragraph(f"Evidence: {ev}")
    steps = info.get("manual_steps")
    if steps:
        doc.add_paragraph("Manual steps:")
        for s in steps:
            doc.add_paragraph(f"- {s}")


def write_docx_catm(path: Path, results: Dict[str, Dict[str, Any]], reqs_by_puid: Dict[str, Dict[str, Any]], template: Optional[Path]) -> None:
    if not DOCX_AVAILABLE:
        return
    doc = Document(template.as_posix()) if (template and template.exists()) else Document()
    if not template:
        doc.add_heading("SEC-CATm — Detailed Catalog", level=1)
    for puid, info in results.items():
        _append_requirement_block(doc, puid, info, reqs_by_puid)
    doc.save(path.as_posix())


def write_docx_summary(path: Path, totals: Dict[str, int], template: Optional[Path]) -> None:
    if not DOCX_AVAILABLE:
        return
    doc = Document(template.as_posix()) if (template and template.exists()) else Document()
    if not template:
        doc.add_heading("Audit Summary", level=1)
        p = doc.add_paragraph(
            f"Total requirements: {totals['total']} | YES: {totals['YES']} | NO: {totals['NO']} | N/A: {totals['N/A']} | MANUAL: {totals['MANUAL']}"
        )
        p.runs[0].font.size = Pt(12)
    else:
        # If there is a table in the template, try to fill the first empty row
        # (We keep the behavior simple/robust; your template already provides the structure)
        doc.add_paragraph(
            f"Total requirements: {totals['total']} | YES: {totals['YES']} | NO: {totals['NO']} | N/A: {totals['N/A']} | MANUAL: {totals['MANUAL']}"
        )
    doc.save(path.as_posix())


# ---- Main --------------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(description="SEC-CATm AI Correlator")
    ap.add_argument("--checklist", type=Path, required=True, help="Path to requirements.json")
    ap.add_argument("--reports", type=Path, required=True, help="Folder containing artifacts (MobSF/SARIF/Trivy)")
    ap.add_argument("--source-root", type=Path, required=True, help="Repository root for quick hints")
    ap.add_argument("--output-dir", type=Path, required=True, help="Output folder")
    ap.add_argument("--max-requirements", type=int, default=-1, help="Limit the number of requirements to audit")
    ap.add_argument("--preselect", action="store_true", help="Preselect src-catm* based on capabilities")
    ap.add_argument("--mark-manual-when-uncertain", action="store_true", help="N/A uncertain -> MANUAL")
    ap.add_argument("--use-llm", choices=["never", "uncertain", "always"], default="never",
                    help="When to invoke the LLM (default: never)")
    ap.add_argument("--llm-timeout", type=int, default=25, help="Soft timeout per requirement (s)")
    ap.add_argument("--app-context", type=Path, default=None, help="File with context (README/docs)")
    ap.add_argument("--catm-template", type=Path, default=None, help="DOCX template for secm-catm")
    ap.add_argument("--summary-template", type=Path, default=None, help="DOCX template for audit summary")
    ap.add_argument("--verbose", action="count", default=0)
    args = ap.parse_args()

    logging.basicConfig(
        level=logging.WARNING - min(args.verbose, 2) * 10,
        format="%(levelname)s | %(message)s",
    )

    outdir = args.output_dir
    outdir.mkdir(parents=True, exist_ok=True)

    # 1) Load
    reqs = load_checklist(args.checklist)
    if args.max_requirements and args.max_requirements > 0:
        reqs = reqs[: args.max_requirements]
    reqs_by_puid: Dict[str, Dict[str, Any]] = {r["PUID"]: r for r in reqs if r.get("PUID")}

    # 2) Evidence
    ev = collect_evidence(args.reports, args.source_root)
    code = quick_code_hints(args.source_root)

    # 3) App context
    app_ctx = summarize_reports_for_ai(ev)
    if args.app_context and args.app_context.exists():
        app_ctx += "\n\n" + read_text_safe(args.app_context, limit=8000)

    # 4) Preselection (src-catm*)
    selected: List[Dict[str, Any]] = list(reqs)
    if args.preselect:
        selected = [r for r in reqs if requirement_seems_related(r, code.caps)]

    total_sel = len(selected)
    results: Dict[str, Dict[str, Any]] = {}

    # 5) Decisions + progress
    for idx, r in enumerate(selected, start=1):
        puid = r.get("PUID") or ""
        if not puid:
            continue

        # Progress
        pct = int(idx * 100 / max(total_sel, 1))
        print(f"[{idx}/{total_sel}] ({pct}%) Auditing {puid} ...", flush=True)

        decision = "MANUAL"
        evidence = ""
        extras: Dict[str, Any] = {}

        # Heuristic first (conservative)
        decision, evidence, uncertain = heuristic_decide(r, code, ev)

        # LLM according to the --use-llm policy
        use_llm = args.use_llm
        need_llm = (use_llm == "always") or (use_llm == "uncertain" and (decision in ("MANUAL", "N/A")))

        if need_llm and OPENAI_AVAILABLE and os.getenv("LLM_API_KEY"):
            llm_dec, llm_ev, llm_ex = llm_decide(puid, r, app_ctx, code, ev, timeout_s=args.llm_timeout)
            # Conservative merge policies:
            if llm_dec == "NO":
                decision, evidence, extras = llm_dec, llm_ev, llm_ex
            elif llm_dec == "YES" and decision not in ("NO",):
                decision, evidence, extras = llm_dec, llm_ev, llm_ex
            elif decision in ("MANUAL", "N/A"):
                # if we were still in doubt, we adopt the LLM output
                decision, evidence, extras = llm_dec, llm_ev, llm_ex
            else:
                extras = {}

        if decision == "N/A" and uncertain and args.mark_manual_when_uncertain:
            decision = "MANUAL"

        info: Dict[str, Any] = {"decision": decision, "evidence": evidence}
        if "rationale" in extras:
            info["rationale"] = extras["rationale"]
        if "manual_steps" in extras:
            info["manual_steps"] = extras["manual_steps"]

        results[puid] = info

    # 6) Totals
    totals = {
        "total": len(reqs),
        "YES": sum(1 for x in results.values() if x["decision"] == "YES"),
        "NO": sum(1 for x in results.values() if x["decision"] == "NO"),
        "N/A": sum(1 for x in results.values() if x["decision"] == "N/A"),
        "MANUAL": sum(1 for x in results.values() if x["decision"] == "MANUAL"),
    }

    # 7) Outputs
    write_audit_findings(outdir / "audit-findings.json", results)
    write_markdown_summary(outdir / "audit-summary.md", totals, notes=[])
    write_docx_checklist(outdir / "checklist.docx", results, reqs_by_puid)
    write_docx_catm(outdir / "secm-catm.docx", results, reqs_by_puid, args.catm_template)
    write_docx_summary(outdir / "audit-summary.docx", totals, args.summary_template)

    # Final log
    print(f"[OK] Generated files in {outdir.as_posix()}")
    for f in ["audit-findings.json", "checklist.docx", "secm-catm.docx", "audit-summary.docx", "audit-summary.md"]:
        fp = outdir / f
        print(("OK" if fp.exists() else "MISSING") + f" -> {fp}")

    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except SystemExit as e:
        raise
    except Exception:
        traceback.print_exc(file=sys.stderr)
        sys.exit(0)
