#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import os
import re
from datetime import datetime, date
from typing import Dict, Any, List
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

LITERALS_PATH = Path(__file__).resolve().parent / "audit_summary_literals.json"
with open(LITERALS_PATH, "r", encoding="utf-8") as f:
    literals = json.load(f)

# Optional LLM (style enhancement)
USE_LLM = True
try:
    from openai import OpenAI
except Exception:
    USE_LLM = False

DEFAULT_IN = "/mnt/data/audit_summary_analysis_pack.json"
DEFAULT_OUT = "/mnt/data/Audit Summary.docx"
CHART_DIR = "/mnt/data/_audit_summary_charts"

HEADER_TEXT = "mSEC-AT Audit Summary - " + literals["header_text"]

COVER_PAGE = literals["cover_page_text"]

def _wrap_label(s: str, width: int = 30) -> str:
    s = str(s)
    if len(s) <= width:
        return s
    out = []
    while len(s) > width:
        cut = s.rfind(" ", 0, width)
        if cut == -1:
            cut = width
        out.append(s[:cut].strip())
        s = s[cut:].strip()
    if s:
        out.append(s)
    return "\n".join(out)

def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _add_field_run(paragraph, field_instr: str) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def _add_header_footer(section, audit_date_str: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = HEADER_TEXT
    if p.runs:
        p.runs[0].font.size = Pt(9)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = f"{audit_date_str} | "
    if fp.runs:
        fp.runs[0].font.size = Pt(9)
    fp.add_run("Page ")
    _add_field_run(fp, " PAGE ")
    fp.add_run(" of ")
    _add_field_run(fp, " NUMPAGES ")
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_cover(doc: Document, audit_date_str: str, auditor: str) -> None:
    doc.add_paragraph()
    t = doc.add_paragraph("mSEC-AT Audit Summary")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28)
    t.runs[0].bold = True

    s = doc.add_paragraph(COVER_PAGE)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(16)
    s.runs[0].bold = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(
        f"Audit Date: {audit_date_str}\n"
        f"Auditor: {auditor}\n"
        f"Classification: Confidential / Internal Use"
    )
    r.font.size = Pt(11)

    doc.add_page_break()

def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.xpath('./w:updateFields')
    if not existing:
        node = OxmlElement("w:updateFields")
        node.set(qn("w:val"), "true")
        settings.append(node)


def _make_bookmark_name(text: str, used: set[str]) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", str(text)).strip("_")
    if not base:
        base = "section"
    if base[0].isdigit():
        base = f"s_{base}"
    name = base[:32]
    seed = name
    i = 2
    while name in used:
        suffix = f"_{i}"
        name = f"{seed[:32-len(suffix)]}{suffix}"
        i += 1
    used.add(name)
    return name


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _render_clickable_toc(after_paragraph, toc_entries: List[tuple[int, str, str]]) -> None:
    current = after_paragraph
    for level, title, anchor in toc_entries:
        p = _insert_paragraph_after(current)
        p.paragraph_format.left_indent = Inches(0.25 * max(level - 1, 0))
        p.paragraph_format.space_after = Pt(2)
        _add_internal_hyperlink(p, title, anchor)
        current = p

def _add_toc(doc: Document):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    doc.add_page_break()
    return p

def _add_manual_toc(doc: Document) -> None:
    doc.add_paragraph("Table of Contents", style="Title")
    items = [
        "1. App information",
        "2. Actors",
        "3. Scope and limitations",
        "4. Evidence criteria",
        "5. Audit summary",
        "6. Main deficiencies",
        "7. Recommendations",
        "8. Visual Analytics",
        "9. Management Action Plan (MAP)",
        "Appendix A — Traceability index (non-exhaustive)",
        "Appendix B — Positive controls verification (workbook traceability)"
    ]
    for item in items:
        doc.add_paragraph(item)
    doc.add_page_break()

def _add_two_col_table(doc: Document, rows: List[List[str]]) -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for k, v in rows:
        cells = tbl.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()


def _add_callout(doc: Document, title: str, bullets: List[str]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, "EDEDED")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    for b in bullets:
        cell.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()


def _add_figure(doc: Document, img_path: str, caption: str) -> None:
    doc.add_picture(img_path, width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _donut(values: List[int], labels: List[str], title: str, center_text: str, out_path: str) -> None:
    fig, ax = plt.subplots(figsize=(7.4, 4.8))
    wedges, _, _ = ax.pie(values, labels=None, autopct="%1.1f%%", startangle=90, pctdistance=0.82)
    centre = plt.Circle((0, 0), 0.55, fc="white")
    fig.gca().add_artist(centre)
    ax.axis("equal")
    ax.set_title(title)
    ax.legend(wedges, labels, loc="center left", bbox_to_anchor=(1.02, 0.5), frameon=False)
    ax.text(0, 0, center_text, ha="center", va="center", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_share_noncompliances(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = []
    for _, d in cat_stats.items():
        items.append((d["category_name"], int(d.get("non_compliant", 0))))
    items = sorted(items, key=lambda x: x[1])

    names = [x[0] for x in items]
    counts = [float(x[1]) for x in items]
    total = float(sum(counts))
    if total <= 0.0:
        pct = [0.0 for _ in counts]
    else:
        pct = [(c / total) * 100.0 for c in counts]

    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct)

    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Share of total non-compliances (%)")
    ax.set_title("Share of non-compliances by category (high-level)")
    ax.set_xlim(0, max(5, float(max(pct) if pct else 0.0) + 5))
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)

    for b, c, p in zip(bars, counts, pct):
        ax.text(b.get_width() + 0.6, b.get_y() + b.get_height() / 2, f"{p:.1f}%  (n={int(c)})", va="center", fontsize=9)

    fig.text(0.01, 0.01, f"Source: audit workbook. Total non-compliances: {int(total)}.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _hbar_compliance_rate(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = []
    for _, d in cat_stats.items():
        app = int(d.get("applicable", 0))
        comp = int(d.get("compliant", 0))
        pct = float(d.get("compliance_pct", 0.0))
        items.append((d["category_name"], pct, comp, app))
    items = sorted(items, key=lambda x: x[1])

    names = [x[0] for x in items]
    pct_vals = [x[1] for x in items]
    comp_vals = [x[2] for x in items]
    app_vals = [x[3] for x in items]

    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))
    bars = ax.barh(y, pct_vals)

    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(n, 30) for n in names])
    ax.set_xlabel("Compliance rate (%) — applicable controls only")
    ax.set_title("Compliance rate by category (applicable controls only)")
    ax.set_xlim(0, 100)
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)

    for b, p, c, a in zip(bars, pct_vals, comp_vals, app_vals):
        ax.text(p + 1, b.get_y() + b.get_height() / 2, f"{p:.1f}%  ({c}/{a})", va="center", fontsize=9)

    fig.text(0.01, 0.01, "Note: (c/a) indicates Compliant / Applicable controls per category. Source: audit workbook.", fontsize=9)
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _stacked_counts(cat_stats: Dict[str, Any], out_path: str) -> None:
    items = []
    for _, d in cat_stats.items():
        items.append((
            d["category_name"],
            int(d.get("compliant", 0)),
            int(d.get("non_compliant", 0)),
            int(d.get("not_applicable", 0)),
        ))
    items = sorted(items, key=lambda x: (x[1] + x[2] + x[3]), reverse=True)

    names = [x[0] for x in items]
    c = [x[1] for x in items]
    n = [x[2] for x in items]
    na = [x[3] for x in items]

    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    y = list(range(len(names)))

    ax.barh(y, c, label="Compliant")
    ax.barh(y, n, left=c, label="Non-compliant")
    left_cn = [cc + nn for cc, nn in zip(c, n)]
    ax.barh(y, na, left=left_cn, label="Not applicable")

    ax.set_yticks(y)
    ax.set_yticklabels([_wrap_label(nm, 30) for nm in names])
    ax.invert_yaxis()
    ax.set_xlabel("Count")
    ax.set_title("Counts by category and status")
    ax.grid(axis="x", linestyle="--", linewidth=0.5, alpha=0.6)
    ax.legend(loc="lower right", frameon=False)

    fig.tight_layout()
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _likelihood_from_count(cnt: int) -> str:
    if cnt >= 50:
        return "High"
    if cnt >= 20:
        return "Medium–High"
    if cnt >= 10:
        return "Medium"
    return "Low–Medium"


def _target_timeline(sev: str) -> str:
    return "0–90 days" if sev == "High" else ("0–180 days" if sev == "Medium" else "0–365 days")


def _target_date_str(audit_dt: date, sev: str) -> str:
    days = 90 if sev == "High" else (180 if sev == "Medium" else 365)
    due = audit_dt.toordinal() + days
    return date.fromordinal(due).strftime("%d %b %Y")


def _extract_json_object(text: str) -> str:
    m = re.search(r"\{.*\}\s*$", text.strip(), flags=re.DOTALL)
    if not m:
        raise ValueError("No JSON object found in model output.")
    return m.group(0)


def _call_llm_for_style(patterns: List[Dict[str, Any]], likelihood_rubric: Dict[str, str], max_takeaways: int = 7) -> Dict[str, Any]:
    """
    Uses LLM only to improve prose and provide richer recommendation bullet sets.
    Does NOT generate any numeric metrics. Counts and rubrics are provided as constraints.
    """
    api_key = os.getenv("LLM_API_KEY", "").strip()
    if not api_key:
        return {}

    model = os.getenv("LLM_MODEL")
    max_tokens = int(os.getenv("LLM_MAX_OUTPUT_TOKENS", "6000"))
    effort = os.getenv("LLM_REASONING_EFFORT", "medium").strip() or "medium"

    #client = OpenAI(api_key=api_key)
    client = OpenAI(
        api_key=api_key,
        base_url=os.getenv("LLM_BASE_URL"),
    )

    inp = []
    for p in patterns[:10]:
        cnt = int(p.get("mapped_noncompliant_count", 0))
        inp.append({
            "pattern": p.get("pattern", ""),
            "count": cnt,
            "severity": p.get("severity", "Low"),
            "likelihood": _likelihood_from_count(cnt),
            "owner": p.get("recommended_owner", "Engineering"),
            "anchors": (p.get("description_anchors", []) or [])[:2],
        })

    system = (
        "You are a senior security audit reporting specialist for peer-review publications. "
        "You will improve wording and generate actionable recommendations ONLY at the weakness-pattern level. "
        "You must NOT invent specific implemented controls. You must NOT invent metrics. "
        "You must NOT claim facts beyond the provided anchors and counts. "
        "Return strict JSON only."
    )

    user_payload = {
        "task": "Generate paper-quality prose components grounded in workbook-derived prevalence counts.",
        "constraints": {
            "no_invented_metrics": True,
            "no_category_level_bullet_dumps": True,
            "no_long_id_lists": True,
            "recommendations_no_time_headings": True,
            "max_key_takeaways": max_takeaways,
            "likelihood_rubric": likelihood_rubric,
        },
        "input_patterns": inp,
        "required_output_schema": {
            "key_takeaways": ["<5-7 bullets; each references prevalence count and pattern name>"],
            "pattern_writeups": [
                {
                    "pattern": "<exact pattern name from input>",
                    "expected": "<1-2 sentences>",
                    "impact": "<1-2 sentences; CIA + regulatory for health data>",
                    "recommendations": ["<6-10 bullets; practical; may include MFA/biometric step-up if appropriate>"]
                }
            ]
        }
    }

    resp = client.responses.create(
        model=model,
        input=[
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
        ],
        max_output_tokens=max_tokens,
        reasoning={"effort": effort},
    )

    txt = (getattr(resp, "output_text", "") or "").strip()
    obj = json.loads(_extract_json_object(txt))
    return obj if isinstance(obj, dict) else {}


def main() -> None:
    in_path = os.getenv("AUDIT_ANALYSIS_JSON_PATH", DEFAULT_IN)
    out_path = os.getenv("AUDIT_SUMMARY_DOCX_PATH", DEFAULT_OUT)

    if not os.path.isfile(in_path):
        raise SystemExit(f"[ERROR] analysis pack not found: {in_path}")

    os.makedirs(CHART_DIR, exist_ok=True)

    with open(in_path, "r", encoding="utf-8") as f:
        pack = json.load(f)

    metrics = pack["metrics"]
    cat_stats = pack["category_metrics"]
    app = pack["app_metadata"]
    actors = pack["actors"]
    patterns = pack["weakness_patterns"]
    pos_controls = pack.get("positive_controls_candidates", [])[:7]
    likelihood_rubric = pack.get("likelihood_rubric", {})

    audit_dt = datetime.utcnow().date()
    audit_date_str = audit_dt.strftime("%d %b %Y")

    plt.rcParams.update({
        "font.size": 10,
        "figure.titlesize": 12,
        "axes.titlesize": 12,
        "axes.labelsize": 10,
    })

    fig1 = os.path.join(CHART_DIR, "figure1_overall_donut.png")
    fig2 = os.path.join(CHART_DIR, "figure2_noncompliance_share_hbar.png")
    fig3 = os.path.join(CHART_DIR, "figure3_compliance_rate_hbar.png")
    fig4 = os.path.join(CHART_DIR, "figure4_counts_stacked_hbar.png")

    applicable = int(metrics["applicable"])
    compliant = int(metrics["compliant"])
    non_compliant = int(metrics["non_compliant"])
    not_applicable = int(metrics["not_applicable"])
    overall_pct = float(metrics["overall_compliance_pct"])

    _donut(
        [compliant, non_compliant, not_applicable],
        ["Compliant", "Non-compliant", "Not applicable"],
        "Overall compliance distribution (workbook-derived)",
        f"{overall_pct:.2f}%\ncompliant\n(applicable)",
        fig1
    )
    _hbar_share_noncompliances(cat_stats, fig2)
    _hbar_compliance_rate(cat_stats, fig3)
    _stacked_counts(cat_stats, fig4)

    prose: Dict[str, Any] = {}
    if USE_LLM and os.getenv("LLM_API_KEY", "").strip():
        try:
            prose = _call_llm_for_style(patterns, likelihood_rubric, max_takeaways=7)
        except Exception:
            prose = {}

    key_takeaways = prose.get("key_takeaways", [])
    writeups = {w["pattern"]: w for w in prose.get("pattern_writeups", []) if isinstance(w, dict) and "pattern" in w}

    if not key_takeaways:
        key_takeaways = []
        for p in patterns[:7]:
            key_takeaways.append(
                f"{p['pattern']} — {p['severity']} severity; {int(p['mapped_noncompliant_count'])} related non-compliant control(s) in the workbook."
            )

    doc = Document()
    _enable_update_fields_on_open(doc)
    _set_doc_defaults(doc)
    _add_header_footer(doc.sections[0], audit_date_str)

    _add_cover(doc, audit_date_str, actors["Auditor"])
    used_bookmarks: set[str] = set()
    toc_entries: List[tuple[int, str, str]] = []
    bookmark_id = 1000

    def add_nav_heading(text: str, level: int):
        nonlocal bookmark_id
        style = f"Heading {level}"
        p = doc.add_paragraph(text, style=style)
        anchor = _make_bookmark_name(text, used_bookmarks)
        _add_bookmark(p, anchor, bookmark_id)
        toc_entries.append((level, text, anchor))
        bookmark_id += 1
        return p

    toc_placeholder = _add_toc(doc)

    add_nav_heading("1. App information", 1)
    _add_two_col_table(doc, [[k, v] for k, v in app.items()])

    add_nav_heading("2. Actors", 1)
    _add_two_col_table(doc, [["Auditor", actors["Auditor"]], ["Requirement Engineering team", "\n".join(actors["Requirement Engineering team"])], ["Engineering Group (EN)", "\n".join(actors["Engineering Group (EN)"])]])

    add_nav_heading("3. Scope and limitations", 1)
    doc.add_paragraph("This Audit Summary consolidates the compliance determinations recorded in the audit workbook. The results reflect the assessed application version and the workbook-defined scope. Controls not evidenced as implemented in the workbook are reported as non-compliant for summary purposes.")

    add_nav_heading("4. Evidence criteria", 1)
    doc.add_paragraph("- Compliant: the workbook provides sufficient evidence that the control is implemented and effective for the assessed scope.\n- Non-compliant: the workbook indicates the control is missing, insufficient, or not evidenced.\n- Not applicable: the control is recorded as out of scope or not relevant for the assessed context.")

    add_nav_heading("5. Audit summary", 1)
    doc.add_paragraph("The audit was carried out using the mSEC-AT (mobile SECurity Audit Tool).")
    doc.add_paragraph(f"Overall, {int(metrics['total_assessed'])} requirements were assessed. {applicable} were applicable controls and {not_applicable} were recorded as not applicable. Of the applicable controls, {compliant} were compliant and {non_compliant} were non-compliant, resulting in an overall compliance rate of {overall_pct:.2f}% (applicable controls only).")
    doc.add_paragraph("This report summarizes the dominant weakness patterns evidenced by non-compliant requirements and proposes actionable remediations suitable for mHealth/EMR environments handling sensitive health information.")

    add_nav_heading("5.1 Key takeaways (Top findings)", 2)
    _add_callout(doc, "Key takeaways (Top findings)", key_takeaways[:7])

    add_nav_heading("5.2 Positive controls observed", 2)
    doc.add_paragraph("All statements below are derived exclusively from controls recorded as Compliant in the audit workbook and include supporting signals (flags and/or evidence). Verification traceability is provided in Appendix B.")
    if pos_controls:
        for pc in pos_controls:
            doc.add_paragraph(pc["declarative_statement"], style="List Bullet")
    else:
        doc.add_paragraph("No compliant controls with supporting evidence/flags were available for verification in the workbook.", style="List Bullet")

    add_nav_heading("5.3 Risk scoring approach", 2)
    doc.add_paragraph("Severity and likelihood ratings in this report follow a qualitative rubric grounded in the audit workbook:\n- Severity reflects potential impact on confidentiality, integrity, and availability of health information, including regulatory exposure.\n- Likelihood is derived from workbook prevalence: the count of non-compliant controls mapped to a weakness pattern as a proxy for exposure.\nLikelihood mapping: High (>=50), Medium-High (20-49), Medium (10-19), Low-Medium (<10).")

    add_nav_heading("5.4 Risk triage (prioritized)", 2)
    rt = doc.add_table(rows=1, cols=7)
    rt.style = "Table Grid"
    h = rt.rows[0].cells
    for idx, txt in enumerate(["Weakness pattern", "Severity (rubric)", "Impact", "Likelihood (workbook prevalence)", "Workbook basis", "Recommended owner", "Target timeline"]):
        h[idx].text = txt
        _set_cell_shading(h[idx], "D9E1F2")
        for run in h[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        cnt = int(p["mapped_noncompliant_count"])
        lik = _likelihood_from_count(cnt)
        sev = p["severity"]
        owner = p["recommended_owner"]
        impact = writeups.get(p["pattern"], {}).get("impact", "The weakness pattern can compromise confidentiality/integrity/availability of health information and increase regulatory exposure.")
        row = rt.add_row().cells
        row[0].text = p["pattern"]
        row[1].text = sev
        row[2].text = impact
        row[3].text = lik
        row[4].text = f"{cnt} mapped non-compliant control(s) in the workbook."
        row[5].text = owner
        row[6].text = _target_timeline(sev)
    doc.add_paragraph()

    add_nav_heading("6. Main deficiencies", 1)
    doc.add_paragraph("The following deficiencies are synthesized as common weakness patterns based on non-compliant requirements. They are not grouped by category; instead they represent cross-cutting gaps evidenced in the audit workbook.")
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        owner = p["recommended_owner"]
        ex_ids = p.get("example_puids", [])[:4]
        anchors = p.get("description_anchors", [])[:2]
        doc.add_paragraph(f"{pat} ({sev})", style="Heading 2")
        doc.add_paragraph(f"Workbook basis: {cnt} related non-compliant control(s) mapped to this pattern.")
        expected = writeups.get(pat, {}).get("expected", "Controls in this area should provide robust, consistently enforced safeguards appropriate to health data processing.")
        impact = writeups.get(pat, {}).get("impact", "Deficiencies can increase the likelihood and impact of security incidents affecting confidentiality, integrity, or availability.")
        doc.add_paragraph(f"Expected: {expected}")
        doc.add_paragraph("Observed: The audit workbook indicates the related controls are missing, insufficient, or not evidenced for the assessed scope.")
        doc.add_paragraph(f"Impact: {impact}")
        doc.add_paragraph(f"Recommended owner: {owner}")
        if ex_ids:
            doc.add_paragraph(f"Traceability (examples, non-exhaustive): {', '.join(ex_ids)}.")
        for a in anchors:
            doc.add_paragraph(f"Evidence anchor (from workbook description): {a}", style="List Bullet")

    doc.add_page_break()
    add_nav_heading("7. Recommendations", 1)
    doc.add_paragraph("Recommendations are organized by the same weakness patterns presented in the Main deficiencies section. They target remediation of workbook-evidenced gaps and may include strengthening controls to improve security posture.")
    for p in patterns[:10]:
        pat = p["pattern"]
        doc.add_paragraph(pat, style="Heading 2")
        recs = writeups.get(pat, {}).get("recommendations", [])
        if not recs:
            recs = [
                "Define and document secure-by-default requirements for this control area, and implement automated tests to prevent regressions.",
                "Apply least-privilege, defense-in-depth, and secure configuration baselines aligned with mHealth/EMR risk profiles.",
                "Introduce step-up authentication, such as MFA, TOTP, FIDO2, or Android BiometricPrompt, for sensitive actions where feasible.",
                "Validate effectiveness through security testing and re-assessment of the mapped non-compliant controls.",
            ]
        for r in recs[:12]:
            doc.add_paragraph(str(r).strip(), style="List Bullet")

    doc.add_page_break()
    add_nav_heading("8. Visual Analytics", 1)
    doc.add_paragraph("Figures below summarize workbook-derived outcomes and distributions. All figures: source: audit workbook.")
    _add_figure(doc, fig1, "Figure 1. Overall compliance distribution (donut chart; source: audit workbook).")
    _add_figure(doc, fig2, "Figure 2. Share of non-compliances by category (legible horizontal bars; source: audit workbook).")
    _add_figure(doc, fig3, "Figure 3. Compliance rate by category (applicable controls only; source: audit workbook).")
    _add_figure(doc, fig4, "Figure 4. Counts by category and status (horizontal stacked bars; source: audit workbook).")

    doc.add_page_break()
    add_nav_heading("9. Management Action Plan (MAP)", 1)
    doc.add_paragraph("Severity and likelihood nomenclature follow the rubric in Section 5.3. Likelihood is supported by workbook prevalence counts recorded in the Workbook basis column.")
    mp = doc.add_table(rows=1, cols=9)
    mp.style = "Table Grid"
    mh = mp.rows[0].cells
    headers = ["Finding / weakness pattern", "Severity", "Likelihood", "Workbook basis", "Owner", "Management action", "Target window", "Target date", "Acceptance criteria / KPI"]
    for idx, txt in enumerate(headers):
        mh[idx].text = txt
        _set_cell_shading(mh[idx], "D9E1F2")
        for run in mh[idx].paragraphs[0].runs:
            run.bold = True
    for p in patterns[:10]:
        pat = p["pattern"]
        cnt = int(p["mapped_noncompliant_count"])
        sev = p["severity"]
        lik = _likelihood_from_count(cnt)
        owner = p["recommended_owner"]
        target_window = _target_timeline(sev)
        target_date = _target_date_str(audit_dt, sev)
        recs = writeups.get(pat, {}).get("recommendations", [])
        action = " ".join([r.strip() for r in recs[:3]]) if recs else "Implement remediation actions aligned to the weakness pattern and validate effectiveness."
        row = mp.add_row().cells
        row[0].text = pat
        row[1].text = sev
        row[2].text = lik
        row[3].text = f"{cnt} mapped non-compliant control(s) in workbook."
        row[4].text = owner
        row[5].text = action
        row[6].text = target_window
        row[7].text = target_date
        row[8].text = "Evidence recorded (tests/configs/release refs); mapped controls can be re-tested and re-scored as compliant."

    doc.add_page_break()
    add_nav_heading("Appendix A - Traceability index (non-exhaustive)", 1)
    doc.add_paragraph("For complete traceability and evidence, refer to the audit workbook.")
    for p in patterns[:10]:
        ex = p.get("example_puids", [])[:5]
        if ex:
            doc.add_paragraph(f"{p['pattern']}: {', '.join(ex)}", style="List Bullet")

    doc.add_page_break()
    add_nav_heading("Appendix B - Positive controls verification (workbook traceability)", 1)
    doc.add_paragraph("This appendix verifies each Positive controls observed statement by providing the originating PUID, flags used, and an evidence excerpt when available.")
    vb = doc.add_table(rows=1, cols=4)
    vb.style = "Table Grid"
    vh = vb.rows[0].cells
    for idx, txt in enumerate(["Positive control statement (as reported)", "Workbook PUID", "Flags used", "Evidence / justification (excerpt)"]):
        vh[idx].text = txt
        _set_cell_shading(vh[idx], "D9E1F2")
        for run in vh[idx].paragraphs[0].runs:
            run.bold = True
    if pos_controls:
        for pc in pos_controls:
            r = vb.add_row().cells
            r[0].text = pc["declarative_statement"]
            r[1].text = pc["puid"]
            r[2].text = pc.get("flags_used", "") or ""
            r[3].text = pc.get("evidence_excerpt", "") or ""
    else:
        r = vb.add_row().cells
        r[0].text = "No verified positive controls available."
        r[1].text = ""
        r[2].text = ""
        r[3].text = ""

    _render_clickable_toc(toc_placeholder, toc_entries)
    _enable_update_fields_on_open(doc)
    doc.save(out_path)
    print(f"[OK] DOCX generated -> {out_path}")


if __name__ == "__main__":
    main()